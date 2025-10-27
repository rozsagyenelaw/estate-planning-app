#!/usr/bin/env python3
"""
Create professional Living Trust template with complete content and POST-PROCESS to fix tags
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import zipfile
import re
from lxml import etree

def create_initial_template():
    """Create template using python-docx (tags will be split)"""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Cover page
    h = doc.add_heading('DECLARATION OF TRUST', 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    h = doc.add_heading('ESTABLISHING THE', 2)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph('{trustName}')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(14)
        run.font.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph('{#if isRestatement}Restatement dated {trustDate}{/if}')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph('{#if notIsRestatement}Dated {trustDate}{/if}')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Article One
    h = doc.add_heading('Article One', 1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h = doc.add_heading('Establishing the Trust', 2)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_paragraph('{#if isRestatement}On {originalTrustDate}, I established the {originalTrustName}, and reserved the right to amend the trust, in whole or in part. On this day, {trustDate}, I revoke all restatements and amendments made to date, and completely restate the trust as follows:{/if}')

    doc.add_paragraph('{#if notIsRestatement}I, {grantorFullName}, also known as the "Grantor," declare that I have set aside and hold in trust all property described in the attached Schedule A. This trust shall be known as the {trustName}.{/if}')

    doc.add_paragraph()

    # Article Two
    h = doc.add_heading('Article Two', 1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h = doc.add_heading('Family Information', 2)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_paragraph('{maritalStatus}')
    doc.add_paragraph()
    doc.add_paragraph('{childrenStatement}')
    doc.add_paragraph()
    doc.add_paragraph('My children are:')
    doc.add_paragraph()
    doc.add_paragraph('{#children}')
    p = doc.add_paragraph('{fullName}, born on {dateOfBirth}', style='List Number')
    doc.add_paragraph('{/children}')
    doc.add_paragraph()

    # Article Three
    h = doc.add_heading('Article Three', 1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h = doc.add_heading('Trust Administration During My Lifetime', 2)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_paragraph('Section 3.1 - Grantor as Trustee')
    doc.add_paragraph('During my lifetime, I shall serve as the Trustee of this trust. As Trustee, I shall have all powers granted under this Declaration of Trust and applicable law.')
    doc.add_paragraph()

    doc.add_paragraph('Section 3.2 - Distribution of Income and Principal')
    doc.add_paragraph('During my lifetime, the Trustee shall distribute to me, or for my benefit, such amounts of the net income and principal of the trust as I may request from time to time.')
    doc.add_paragraph()

    doc.add_paragraph('Section 3.3 - Revocation and Amendment')
    doc.add_paragraph('I reserve the right to revoke or amend this trust, in whole or in part, at any time during my lifetime by delivering a written instrument to the Trustee.')
    doc.add_paragraph()

    # Article Four
    h = doc.add_heading('Article Four', 1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h = doc.add_heading('Successor Trustees', 2)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_paragraph('Section 4.1 - Successor Trustees During Incapacity')
    doc.add_paragraph('If I become incapacitated and unable to serve as Trustee, the following persons shall serve as Successor Trustees, in the order listed:')
    doc.add_paragraph()
    doc.add_paragraph('{successorTrusteesDuringIncapacityFormatted}')
    doc.add_paragraph()

    doc.add_paragraph('Section 4.2 - Successor Trustees After Death')
    doc.add_paragraph('Upon my death, the following persons shall serve as Successor Trustees, in the order listed:')
    doc.add_paragraph()
    doc.add_paragraph('{successorTrusteesAfterDeathFormatted}')
    doc.add_paragraph()

    doc.add_paragraph('Section 4.3 - Trustee Powers')
    doc.add_paragraph('The Successor Trustee shall have all powers necessary to administer this trust.')
    doc.add_paragraph()

    # Article Five
    h = doc.add_heading('Article Five', 1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h = doc.add_heading('Distribution Upon My Death', 2)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_paragraph('Section 5.1 - Specific Distributions')
    doc.add_paragraph('{#if hasSpecificDistributions}Upon my death, the Trustee shall distribute the following specific items of property:{/if}')
    doc.add_paragraph()
    doc.add_paragraph('{#specificDistributions}')
    p = doc.add_paragraph('To {beneficiaryName}: {property}', style='List Bullet')
    doc.add_paragraph('{#if hasAgeCondition}If {beneficiaryName} has not reached age {conditionAge} at the time of distribution, this property shall be held in trust until that age is reached.{/if}')
    doc.add_paragraph('{/specificDistributions}')
    doc.add_paragraph('{#if notHasSpecificDistributions}There are no specific distributions.{/if}')
    doc.add_paragraph()

    doc.add_paragraph('Section 5.2 - Residuary Trust Estate')
    doc.add_paragraph('After payment of debts and expenses, the Trustee shall distribute the remaining trust estate to the following beneficiaries:')
    doc.add_paragraph()
    doc.add_paragraph('{#beneficiaries}')
    p = doc.add_paragraph('{percentage}% to {fullName} {#if dateOfBirth}(born {dateOfBirth}, {relationship}){/if}', style='List Bullet')
    doc.add_paragraph('{/beneficiaries}')
    doc.add_paragraph()

    # Article Six
    h = doc.add_heading('Article Six', 1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h = doc.add_heading('General Needs Trusts', 2)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_paragraph('{#if hasGeneralNeeds}I direct the Trustee to establish the following General Needs Trusts:{/if}')
    doc.add_paragraph()
    doc.add_paragraph('{#generalNeeds}')
    doc.add_paragraph('Section {sectionNumber} - Trust for {beneficiaryName}')
    doc.add_paragraph('The Trustee shall hold the share distributable to {beneficiaryName} in trust for their benefit.')
    doc.add_paragraph('{#if hasAgeCondition}This trust shall terminate when {beneficiaryName} reaches age {terminationAge}.{/if}')
    doc.add_paragraph('{/generalNeeds}')
    doc.add_paragraph('{#if notHasGeneralNeeds}No General Needs Trusts are established.{/if}')
    doc.add_paragraph()

    # Article Seven
    h = doc.add_heading('Article Seven', 1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h = doc.add_heading('Trust Administration Provisions', 2)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_paragraph('Section 7.1 - Payment of Expenses')
    doc.add_paragraph('The Trustee shall pay all debts and expenses from the trust property.')
    doc.add_paragraph()

    doc.add_paragraph('Section 7.2 - No Bond Required')
    doc.add_paragraph('No Successor Trustee shall be required to post bond.')
    doc.add_paragraph()

    doc.add_paragraph('Section 7.3 - Governing Law')
    doc.add_paragraph('This trust shall be governed by applicable state law.')
    doc.add_paragraph()

    # Signature
    doc.add_page_break()
    doc.add_paragraph('IN WITNESS WHEREOF, I have executed this Declaration of Trust on {trustDate}.')
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('________________________________________')
    doc.add_paragraph('{grantorFullName}, Grantor and Trustee')
    doc.add_paragraph()

    # Schedule A
    doc.add_page_break()
    h = doc.add_heading('SCHEDULE A', 1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h = doc.add_heading('Trust Property', 2)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_paragraph('The following property is held in the {trustName}:')
    doc.add_paragraph()
    doc.add_paragraph('{assets}')
    doc.add_paragraph()

    temp_path = '/tmp/template_before_fix.docx'
    doc.save(temp_path)
    print(f"✓ Created initial template (tags will be split)")
    return temp_path

def fix_template_tags(input_path, output_path):
    """Post-process DOCX to merge all template tags into single <w:t> elements"""
    print("\n✓ Post-processing to fix template tags...")

    NAMESPACES = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    }

    # Read DOCX
    with zipfile.ZipFile(input_path, 'r') as zip_in:
        xml_bytes = zip_in.read('word/document.xml')

    # Parse XML
    root = etree.fromstring(xml_bytes)

    # Process each paragraph
    paragraphs = root.findall('.//w:p', NAMESPACES)
    print(f"  Processing {len(paragraphs)} paragraphs...")

    fixed_count = 0
    for p in paragraphs:
        # Get all text from this paragraph
        texts = []
        for t_elem in p.findall('.//w:t', NAMESPACES):
            if t_elem.text:
                texts.append(t_elem.text)

        combined_text = ''.join(texts)

        # Check if this paragraph contains template tags
        if '{' in combined_text and '}' in combined_text:
            # Remove all existing <w:r> elements
            for r_elem in p.findall('.//w:r', NAMESPACES):
                r_elem.getparent().remove(r_elem)

            # Create a single new run with all text
            r = etree.SubElement(p, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
            t = etree.SubElement(r, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t.text = combined_text
            fixed_count += 1

    print(f"  Fixed {fixed_count} paragraphs with template tags")

    # Write back
    fixed_xml = etree.tostring(root, encoding='utf-8', xml_declaration=True)

    with zipfile.ZipFile(input_path, 'r') as zip_in:
        with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zip_out:
            for item in zip_in.infolist():
                if item.filename == 'word/document.xml':
                    zip_out.writestr(item, fixed_xml)
                else:
                    zip_out.writestr(item, zip_in.read(item.filename))

    print(f"✓ Saved fixed template to: {output_path}")

# Main execution
print("Creating professional Living Trust template...\n")
temp_path = create_initial_template()
output_path = './public/templates/single_living_trust_template.docx'
fix_template_tags(temp_path, output_path)

print("\n" + "="*60)
print("✓✓✓ TEMPLATE CREATED SUCCESSFULLY ✓✓✓")
print("="*60)
print(f"\nTemplate saved to: {output_path}")
print("\nThis template includes:")
print("  ✓ Professional formatting with Times New Roman font")
print("  ✓ Conditional logic for restatements")
print("  ✓ Conditional logic for specific distributions")
print("  ✓ Conditional logic for general needs trusts")
print("  ✓ Support for multiple children")
print("  ✓ Support for multiple beneficiaries")
print("  ✓ Support for multiple successor trustees")
print("  ✓ Support for multiple specific distributions")
print("  ✓ Support for multiple general needs trusts")
print("  ✓ All template tags properly formatted (NOT split!)")
print("\nReady to test!")
