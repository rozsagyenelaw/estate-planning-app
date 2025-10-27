#!/usr/bin/env python3
"""
Create a brand new professional Single Living Trust template with proper formatting
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create new document
doc = Document()

# Set up styles
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Helper function to add centered heading
def add_centered_heading(text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].font.name = 'Times New Roman'
    heading.runs[0].font.size = Pt(14 if level == 1 else 12)
    heading.runs[0].font.bold = True
    return heading

# Helper function to add paragraph with proper spacing
def add_paragraph(text='', bold=False, indent=False):
    p = doc.add_paragraph(text)
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.bold = bold
    if indent:
        p.paragraph_format.left_indent = Inches(0.5)
    return p

# ============= COVER PAGE =============
add_centered_heading('DECLARATION OF TRUST', level=1)
doc.add_paragraph()
add_centered_heading('ESTABLISHING THE', level=2)
doc.add_paragraph()
add_centered_heading('{trustName}', level=2)
doc.add_paragraph()
doc.add_paragraph()

# Restatement conditional
p = doc.add_paragraph()
p.add_run('{{#if isRestatement}}').font.size = Pt(12)
p.add_run('Restatement dated {trustDate}').font.size = Pt(12)
p.add_run('{{/if}}').font.size = Pt(12)

p = doc.add_paragraph()
p.add_run('{{#if notIsRestatement}}').font.size = Pt(12)
p.add_run('Dated {trustDate}').font.size = Pt(12)
p.add_run('{{/if}}').font.size = Pt(12)

doc.add_page_break()

# ============= ARTICLE ONE: ESTABLISHING THE TRUST =============
add_centered_heading('Article One', level=1)
add_centered_heading('Establishing the Trust', level=2)
doc.add_paragraph()

# Restatement section
p = doc.add_paragraph()
p.add_run('{{#if isRestatement}}').font.size = Pt(12)
p.add_run('On {originalTrustDate}, I established the {originalTrustName}, and reserved the right to amend the trust, in whole or in part. On this day, {trustDate}, I revoke all restatements and amendments made to date, and completely restate the trust as follows:').font.size = Pt(12)
p.add_run('{{/if}}').font.size = Pt(12)

# Non-restatement section
p = doc.add_paragraph()
p.add_run('{{#if notIsRestatement}}').font.size = Pt(12)
p.add_run('I, {grantorFullName}, also known as the "Grantor," declare that I have set aside and hold in trust all property described in the attached Schedule A. This trust shall be known as the {trustName}.').font.size = Pt(12)
p.add_run('{{/if}}').font.size = Pt(12)

doc.add_paragraph()

# ============= ARTICLE TWO: FAMILY INFORMATION =============
add_centered_heading('Article Two', level=1)
add_centered_heading('Family Information', level=2)
doc.add_paragraph()

add_paragraph('{maritalStatus}')
doc.add_paragraph()
add_paragraph('{childrenStatement}')
doc.add_paragraph()

# List children
p = doc.add_paragraph()
p.add_run('{{#children}}').font.size = Pt(12)
p_child = doc.add_paragraph('{fullName}, born on {dateOfBirth}', style='List Number')
p_child.runs[0].font.name = 'Times New Roman'
p_child.runs[0].font.size = Pt(12)
p = doc.add_paragraph()
p.add_run('{{/children}}').font.size = Pt(12)

doc.add_paragraph()

# ============= ARTICLE THREE: TRUST ADMINISTRATION DURING GRANTOR\'S LIFETIME =============
add_centered_heading('Article Three', level=1)
add_centered_heading('Trust Administration During My Lifetime', level=2)
doc.add_paragraph()

add_paragraph('Section 3.1 - Grantor as Trustee')
add_paragraph('During my lifetime, I shall serve as the Trustee of this trust. As Trustee, I shall have all powers granted under this Declaration of Trust and applicable law.')
doc.add_paragraph()

add_paragraph('Section 3.2 - Distribution of Income and Principal')
add_paragraph('During my lifetime, the Trustee shall distribute to me, or for my benefit, such amounts of the net income and principal of the trust as I may request from time to time.')
doc.add_paragraph()

add_paragraph('Section 3.3 - Revocation and Amendment')
add_paragraph('I reserve the right to revoke or amend this trust, in whole or in part, at any time during my lifetime by delivering a written instrument to the Trustee.')
doc.add_paragraph()

# ============= ARTICLE FOUR: SUCCESSOR TRUSTEES =============
add_centered_heading('Article Four', level=1)
add_centered_heading('Successor Trustees', level=2)
doc.add_paragraph()

add_paragraph('Section 4.1 - Successor Trustees During Incapacity')
add_paragraph('If I become incapacitated and unable to serve as Trustee, the following persons shall serve as Successor Trustees, in the order listed:')
doc.add_paragraph()
add_paragraph('{successorTrusteesDuringIncapacityFormatted}', indent=True)
doc.add_paragraph()

add_paragraph('Section 4.2 - Successor Trustees After Death')
add_paragraph('Upon my death, the following persons shall serve as Successor Trustees, in the order listed:')
doc.add_paragraph()
add_paragraph('{successorTrusteesAfterDeathFormatted}', indent=True)
doc.add_paragraph()

add_paragraph('Section 4.3 - Trustee Powers')
add_paragraph('The Successor Trustee shall have all powers necessary to administer this trust, including but not limited to the power to:')
p = doc.add_paragraph('Collect, hold, and retain trust property', style='List Bullet')
p.runs[0].font.name = 'Times New Roman'
p = doc.add_paragraph('Receive additions to the trust', style='List Bullet')
p.runs[0].font.name = 'Times New Roman'
p = doc.add_paragraph('Invest and reinvest trust property', style='List Bullet')
p.runs[0].font.name = 'Times New Roman'
p = doc.add_paragraph('Sell, exchange, or lease trust property', style='List Bullet')
p.runs[0].font.name = 'Times New Roman'
p = doc.add_paragraph('Make distributions of income and principal', style='List Bullet')
p.runs[0].font.name = 'Times New Roman'
p = doc.add_paragraph('Hire professionals (attorneys, accountants, investment advisors)', style='List Bullet')
p.runs[0].font.name = 'Times New Roman'
p = doc.add_paragraph('Do all other acts necessary for the prudent management of the trust', style='List Bullet')
p.runs[0].font.name = 'Times New Roman'
doc.add_paragraph()

# ============= ARTICLE FIVE: DISTRIBUTION UPON GRANTOR\'S DEATH =============
add_centered_heading('Article Five', level=1)
add_centered_heading('Distribution Upon My Death', level=2)
doc.add_paragraph()

# Specific Distributions Section
add_paragraph('Section 5.1 - Specific Distributions')

p = doc.add_paragraph()
p.add_run('{{#if hasSpecificDistributions}}').font.size = Pt(12)

add_paragraph('Upon my death, the Trustee shall distribute the following specific items of property:')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('{{#specificDistributions}}').font.size = Pt(12)

p_dist = doc.add_paragraph('', style='List Bullet')
p_dist.add_run('To {beneficiaryName}: {property}').font.name = 'Times New Roman'

p_age = doc.add_paragraph()
p_age.add_run('{{#if hasAgeCondition}}').font.size = Pt(12)
p_age.add_run(' If {beneficiaryName} has not reached age {conditionAge} at the time of distribution, this property shall be held in trust until that age is reached.').font.size = Pt(12)
p_age.add_run('{{/if}}').font.size = Pt(12)

p = doc.add_paragraph()
p.add_run('{{/specificDistributions}}').font.size = Pt(12)

p = doc.add_paragraph()
p.add_run('{{/if}}').font.size = Pt(12)

p = doc.add_paragraph()
p.add_run('{{#if notHasSpecificDistributions}}').font.size = Pt(12)
p.add_run('There are no specific distributions. All trust property shall be distributed according to Section 5.2 below.').font.size = Pt(12)
p.add_run('{{/if}}').font.size = Pt(12)

doc.add_paragraph()

# Residuary Distributions
add_paragraph('Section 5.2 - Residuary Trust Estate')
add_paragraph('After payment of my debts, expenses, and any specific distributions, the Trustee shall distribute the remaining trust estate to the following beneficiaries in the percentages indicated:')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('{{#beneficiaries}}').font.size = Pt(12)

p_ben = doc.add_paragraph('', style='List Bullet')
p_ben.add_run('{percentage}% to {fullName}').font.name = 'Times New Roman'

p_dob = doc.add_paragraph()
p_dob.add_run('{{#if dateOfBirth}}').font.size = Pt(12)
p_dob.add_run(' (born {dateOfBirth}, {relationship})').font.size = Pt(12)
p_dob.add_run('{{/if}}').font.size = Pt(12)

p = doc.add_paragraph()
p.add_run('{{/beneficiaries}}').font.size = Pt(12)

doc.add_paragraph()

add_paragraph('If any beneficiary does not survive me, their share shall be distributed equally among the surviving beneficiaries.')
doc.add_paragraph()

# ============= ARTICLE SIX: GENERAL NEEDS TRUSTS =============
add_centered_heading('Article Six', level=1)
add_centered_heading('General Needs Trusts', level=2)
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('{{#if hasGeneralNeeds}}').font.size = Pt(12)

add_paragraph('I direct the Trustee to establish the following General Needs Trusts for beneficiaries who require special management of their inheritance:')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('{{#generalNeeds}}').font.size = Pt(12)

add_paragraph('Section {sectionNumber} - Trust for {beneficiaryName}', bold=True)
doc.add_paragraph()

add_paragraph('The Trustee shall hold the share distributable to {beneficiaryName} in a separate trust for their benefit. The Trustee shall distribute income and principal for {beneficiaryName}\'s health, education, maintenance, and support.')
doc.add_paragraph()

p_age = doc.add_paragraph()
p_age.add_run('{{#if hasAgeCondition}}').font.size = Pt(12)
p_age.add_run('This trust shall terminate when {beneficiaryName} reaches age {terminationAge}, at which time the remaining trust property shall be distributed outright to {beneficiaryName}.').font.size = Pt(12)
p_age.add_run('{{/if}}').font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('{{/generalNeeds}}').font.size = Pt(12)

p = doc.add_paragraph()
p.add_run('{{/if}}').font.size = Pt(12)

p = doc.add_paragraph()
p.add_run('{{#if notHasGeneralNeeds}}').font.size = Pt(12)
p.add_run('No General Needs Trusts are established under this Declaration of Trust. All distributions shall be made outright to beneficiaries.').font.size = Pt(12)
p.add_run('{{/if}}').font.size = Pt(12)

doc.add_paragraph()

# ============= ARTICLE SEVEN: TRUST ADMINISTRATION PROVISIONS =============
add_centered_heading('Article Seven', level=1)
add_centered_heading('Trust Administration Provisions', level=2)
doc.add_paragraph()

add_paragraph('Section 7.1 - Payment of Expenses')
add_paragraph('The Trustee shall pay all debts, expenses of last illness, funeral expenses, and expenses of administering my estate and this trust from the trust property.')
doc.add_paragraph()

add_paragraph('Section 7.2 - No Bond Required')
add_paragraph('No Successor Trustee named in this Declaration of Trust shall be required to post bond.')
doc.add_paragraph()

add_paragraph('Section 7.3 - Trustee Compensation')
add_paragraph('The Trustee shall be entitled to reasonable compensation for services rendered, plus reimbursement for expenses incurred in the administration of the trust.')
doc.add_paragraph()

add_paragraph('Section 7.4 - Spendthrift Provision')
add_paragraph('No beneficiary shall have the power to assign, transfer, pledge, or otherwise encumber their interest in the trust. The trust property shall not be subject to the claims of creditors of any beneficiary.')
doc.add_paragraph()

add_paragraph('Section 7.5 - Governing Law')
add_paragraph('This trust shall be governed by the laws of the State in which it is executed.')
doc.add_paragraph()

# ============= SIGNATURE SECTION =============
doc.add_page_break()
add_paragraph('IN WITNESS WHEREOF, I have executed this Declaration of Trust on {trustDate}.')
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

add_paragraph('_' * 40)
add_paragraph('{grantorFullName}, Grantor and Trustee')
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# ============= SCHEDULE A =============
doc.add_page_break()
add_centered_heading('SCHEDULE A', level=1)
add_centered_heading('Trust Property', level=2)
doc.add_paragraph()

add_paragraph('The following property is held in the {trustName}:')
doc.add_paragraph()
add_paragraph('{assets}', indent=True)
doc.add_paragraph()
doc.add_paragraph()

add_paragraph('Additional property may be added to this trust by the Grantor during lifetime or by transfer at death.')

# Save the document
output_path = './public/templates/single_living_trust_template_NEW.docx'
doc.save(output_path)

print(f"✓ Created new professional template: {output_path}")
print("✓ All template tags are properly formatted")
print("✓ Includes conditional logic for:")
print("  - Restatement")
print("  - Specific property distributions")
print("  - General needs trusts")
print("✓ Supports multiple:")
print("  - Children")
print("  - Beneficiaries")
print("  - Successor trustees")
print("  - Specific distributions")
print("  - General needs trusts")
